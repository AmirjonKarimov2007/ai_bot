from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import html
import asyncio


def clean_text(txt: str):
    """Matndan HTML entity-lar va nosoz apostroflarni tozalaydi"""
    if not txt:
        return txt

    # HTML entity-larni oddiy belgiga aylantirish
    txt = html.unescape(txt)

    # Turli ko‘rinishdagi apostroflarni bitta `'` ga almashtirish
    txt = re.sub(r"[‘’´`]", "'", txt)

    # Ikkita yoki ko‘p apostroflarni bitta `'` qilish ('' → ')
    txt = re.sub(r"'+", "'", txt)

    # Ortiqcha bo'shliqlarni tozalash
    return txt.strip()


def add_page_border(doc):
    """Word hujjatiga chiroyli qalin ramka qo‘shadi."""
    sectPr = doc.sections[0]._sectPr

    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')

    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')    # chiziq turi
        border.set(qn('w:sz'), '24')         # qalinlik (24 = qalin)
        border.set(qn('w:space'), '24')      # matnga masofa
        border.set(qn('w:color'), '000000')  # qora rang
        pgBorders.append(border)

    sectPr.append(pgBorders)


async def word_generator(type, mavzu, univer, name, user_id, rejalar: list, theme_text: dict):

    # ----------------------------
    # 1) Kiruvchi matnlarni tozalash
    # ----------------------------
    type = clean_text(type)
    mavzu = clean_text(mavzu)
    univer = clean_text(univer)
    name = clean_text(name)

    doc = Document("Referal.docx")

    # ----------------------------
    # RAMKA QO‘SHISH
    # ----------------------------
    add_page_border(doc)

    # ----------------------------
    # 2) Placeholderlarni almashtirish
    # ----------------------------
    for paragraph in doc.paragraphs:
        old_text = paragraph.text

        new_text = (
            old_text.replace("TypeType", type.upper())
                    .replace("UNIVER,UNIVER,UNIVER,UNIVER", univer.upper())
                    .replace("ThemeTheme", mavzu.title())
                    .replace("namenamename", name.title())
        )

        if old_text != new_text:
            paragraph.clear()
            run = paragraph.add_run(new_text)
            run.font.name = "Times New Roman"

            if "TypeType" in old_text:
                run.font.size = Pt(54)
                run.bold = True
            else:
                run.font.size = Pt(14)

    # ----------------------------
    # 3) Sahifa ajratuvchi
    # ----------------------------
    doc.add_page_break()

    # ----------------------------
    # 4) Rejalar bo‘limi
    # ----------------------------
    heading = doc.add_paragraph("Rejalar")
    hrun = heading.runs[0]
    hrun.font.name = "Times New Roman"
    hrun.font.bold = True
    hrun.font.size = Pt(14)

    for reja in rejalar:
        text = clean_text(reja)
        p = doc.add_paragraph(text)
        run = p.runs[0]
        run.font.name = "Times New Roman"
        run.italic = True
        run.font.size = Pt(14)

    doc.add_page_break()

    # ----------------------------
    # 5) Mavzular bo‘limi
    # ----------------------------
    last_key = list(theme_text.keys())[-1]

    for title, text in theme_text.items():
        title = clean_text(title)
        text = clean_text(text)

        p_title = doc.add_paragraph(title)
        p_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p_title.runs:
            run.font.name = "Times New Roman"
            run.bold = True
            run.font.size = Pt(14)

        p_text = doc.add_paragraph("     " + text)
        p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p_text.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

        if title != last_key:
            doc.add_page_break()

    # ----------------------------
    # 6) Word faylni BytesIO ga yozish
    # ----------------------------
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream
