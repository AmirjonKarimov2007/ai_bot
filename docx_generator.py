import asyncio
from io import BytesIO
import html
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_border_to_section(section):
    """Bitta section (sahifa) ga page border qo'shadi"""
    sectPr = section._sectPr
    
    # Avvalgi borderlarni o'chiramiz (agar bor bo'lsa)
    for borders in sectPr.xpath('./w:pgBorders'):
        sectPr.remove(borders)

    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')

    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')        # qalinligi (12 = 1.5 pt)
        border.set(qn('w:space'), '24')     # chetidan masofa (24 pt)
        border.set(qn('w:color'), '000000') # qora rang
        pgBorders.append(border)

    sectPr.append(pgBorders)

async def word_generator(type, mavzu, univer, name, user_id, rejalar: list, theme_text: dict):
    name = html.unescape(name)
    univer = html.unescape(univer)

    loop = asyncio.get_event_loop()
    doc = await loop.run_in_executor(None, Document, "Referal.docx")

    # 1-sahifa (title page) ga ramka qo'shamiz
    add_page_border_to_section(doc.sections[0])

    # DEBUG: placeholderlarni tekshirish
    print("Placeholderlarni tekshirish...")
    for i, paragraph in enumerate(doc.paragraphs):
        print(f"Paragraph {i}: {paragraph.text}")
        for j, run in enumerate(paragraph.runs):
            print(f"  Run {j}: '{run.text}'")

    # 1. USUL: Avval butun paragraph matnini tekshirish
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "TypeType" in run.text:
                run.text = run.text.replace("TypeType", type.upper())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(54)
                run.bold = True
            if "UNIVER,UNIVER,UNIVER,UNIVER" in run.text:
                run.text = run.text.replace("UNIVER,UNIVER,UNIVER,UNIVER", univer.upper())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
            if "ThemeTheme" in run.text:
                run.text = run.text.replace("ThemeTheme", mavzu.title())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            if "namenamename" in run.text:
                run.text = run.text.replace("namenamename", name.title())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

    

    # 2-sahifa uchun yangi section yaratamiz
    doc.add_page_break()
    doc.add_section()  # Yangi sahifa va yangi section
    add_page_border_to_section(doc.sections[-1])  # Yangi sahifaga ramka

    # Rejalar bo'limi
    p_title = doc.add_paragraph("REJALAR", style='Normal')
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_title.runs:
        run.font.name = 'Times New Roman'
        run.font.bold = True
        run.font.size = Pt(16)

    for reja in rejalar:
        p = doc.add_paragraph(reja, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

    # Har bir mavzu uchun yangi sahifa + yangi section + ramka
    for i, (mavzu_key, text) in enumerate(theme_text.items()):
        doc.add_page_break()
        doc.add_section()  # Yangi sahifa va yangi section
        add_page_border_to_section(doc.sections[-1])  # Har bir yangi sahifaga ramka

        # Mavzu sarlavhasi
        p = doc.add_paragraph(mavzu_key, style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.bold = True
            run.font.size = Pt(16)

        # Matn
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p2.add_run(f"     {text}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # Faylni saqlash
    file_stream = BytesIO()
    await loop.run_in_executor(None, doc.save, file_stream)
    file_stream.seek(0)
    return file_stream

